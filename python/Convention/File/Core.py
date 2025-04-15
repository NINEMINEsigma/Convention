from ..Internal         import *
import                         json
import                         shutil
import pandas           as     pd
import                         os
import                         sys
import                         pickle
import                         zipfile
import                         tarfile
import                         base64
import                         hashlib
import                         time
import                         datetime
import                         stat
from typing             import *
from pathlib            import Path
try:
    from pydub              import AudioSegment
except ImportError:
    ImportingThrow("File", ["pydub"])
try:
    from PIL                import Image, ImageFile
except ImportError:
    ImportingThrow("File", ["Pillow"])
try:
    from docx               import Document
    from docx.document      import Document as DocumentObject
except ImportError:
    ImportingThrow("File", ["python-docx"])

from ..Internal         import *
from ..Str.Core         import UnWrapper, list_byte_to_string
from ..Lang.Core        import split_elements as tool_split_elements

text_readable_file_type = ["txt", "md", "json", "csv", "xml", "xlsx", "xls", "docx", "doc", "svg"]
audio_file_type = ["mp3","ogg","wav"]
image_file_type = ['png', 'jpg', 'jpeg', 'bmp', 'svg', 'ico']
temp_tool_file_path_name = "temp.tool_file"

is_binary_file_functional_test_length:int = 1024

def is_binary_file(file_path: str) -> bool:
    try:
        global is_binary_file_functional_test_length
        with open(file_path, 'rb') as f:
            chunk = f.read(is_binary_file_functional_test_length)  # 读取文件的前缀字节
            return b'\x00' in chunk  # 如果包含NUL字符，则认为是二进制文件
    except Exception as e:
        print(f"Error: {e}")
        return False

def get_extension_name(file:str):
        return os.path.splitext(file)[1][1:]

def get_base_filename(file:str):
    return os.path.basename(file)

def is_image_file(file_path:str):
    return get_extension_name(file_path) in image_file_type

dir_name_type = str
file_name_type = str

class FileOperationError(Exception):
    """文件操作异常基类"""
    pass

class CompressionError(FileOperationError):
    """压缩操作异常"""
    pass

class EncryptionError(FileOperationError):
    """加密操作异常"""
    pass

class HashError(FileOperationError):
    """哈希计算异常"""
    pass

class FileMonitorError(FileOperationError):
    """文件监控异常"""
    pass

class BackupError(FileOperationError):
    """备份操作异常"""
    pass

class PermissionError(FileOperationError):
    """权限操作异常"""
    pass

try:
    from pydantic import BaseModel, GetCoreSchemaHandler
    from pydantic_core import core_schema
except ImportError:
    type BaseModel = Any
    type GetCoreSchemaHandler = Any
    type core_schema = Any

class tool_file(any_class):

    __datas_lit_key:    Literal["model"]    = "model"
    __file_path:        str                 = None
    __file:             IO[Any]             = None
    data:               Any                 = None
    datas:              Dict[str, Any]      = {}


    @classmethod
    def __get_pydantic_core_schema__(
        cls,
        _source_type: Any,
        _handler: GetCoreSchemaHandler,
    ) -> core_schema.CoreSchema:
        return core_schema.no_info_after_validator_function(
            cls,
            core_schema.any_schema(),
            serialization=core_schema.plain_serializer_function_ser_schema(
                lambda instance: None
            ),
        )


    def __init__(
        self,
        file_path:          Union[str, Self],
        file_mode:          str = None,
        *args, **kwargs
        ):
        self.__file_path:   str             = file_path.fullpath if isinstance(file_path, tool_file) else file_path
        self.data:          Any             = None
        self.datas:         Dict[str, Any]  = {}
        self.__file:        IO[Any]         = None
        if file_mode is not None:
            self.open(file_mode, *args, **kwargs)
    def __del__(self):
        self.close()
    def __str__(self):
        return self.get_path()
    def __setitem__(self, key:str, value):
        self.datas[key] = value
    def __getitem__(self, key:str):
        if key not in self.datas:
            self.datas[key] = None
        return self.datas[key]
    def __contains__(self, key:str):
        return key in self.datas
    def __delitem__(self, key:str):
        del self.datas[key]
    def __iter__(self):
        return iter(self.datas)
    def __len__(self):
        return len(self.datas)
    @override
    def __enter__(self):
        if self.is_open():
            return self
        if self.exists() and self.is_file():
            self.load()
        return self
    @override
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()
        return super().__exit__(exc_type, exc_val, exc_tb)

    def __or__(self, other):
        if other is None:
            return tool_file(self.get_path() if self.is_dir() else self.get_path()+"\\")
        else:
            return tool_file(os.path.join(self.get_path(), UnWrapper(other)))
    def __idiv__(self, other):
        self.close()
        temp = self.__or__(other)
        self.__file_path = temp.get_path()

    def to_path(self):
        return Path(self.__file_path)
    def __Path__(self):
        return Path(self.__file_path)

    def write(self, data:Union[str, bytes]):
        self.__file.write(data)

    def create(self):
        if self.exists() == False:
            if self.is_dir():
                if os.path.exists(self.get_dir()):
                    os.makedirs(self.__file_path)
                else:
                    raise FileNotFoundError(f"{self.__file_path} cannt create, because its parent path is not exist")
            else:
                self.open('w')
                self.close()
        return self
    def exists(self):
        return os.path.exists(self.__file_path)
    def remove(self):
        self.close()
        if self.exists():
            if self.is_dir():
                shutil.rmtree(self.__file_path)
            else:
                os.remove(self.__file_path)
        return self
    def copy(self, to_path:Optional[Union[Self, str]]=None):
        if to_path is None:
            return tool_file(self.__file_path)
        if self.exists() is False:
            raise FileNotFoundError("file not found")
        self.close()
        target_file = tool_file(UnWrapper(to_path))
        if target_file.is_dir():
            target_file = target_file|self.get_filename()
        shutil.copy(self.__file_path, UnWrapper(target_file))
        return target_file
    def move(self, to_path:Union[Self, str]):
        if self.exists() is False:
            raise FileNotFoundError("file not found")
        self.close()
        target_file = tool_file(UnWrapper(to_path))
        if target_file.is_dir():
            target_file = target_file|self.get_filename()
        shutil.move(self.__file_path, UnWrapper(target_file))
        self.__file_path = target_file.__file_path
        return self
    def rename(self, newpath:Union[Self, str]):
        if self.exists() is False:
            raise FileNotFoundError("file not found")
        self.close()
        newpath:str = UnWrapper(newpath)
        if '\\' in newpath or '/' in newpath:
            newpath = get_base_filename(newpath)
        new_current_path = os.path.join(self.get_dir(), newpath)
        os.rename(self.__file_path, new_current_path)
        self.__file_path = new_current_path
        return self

    def refresh(self):
        self.load()
        return self
    def open(self, mode='r', is_refresh=False, encoding:str='utf-8', *args, **kwargs):
        self.close()
        if 'b' in mode:
            self.__file = open(self.__file_path, mode, *args, **kwargs)
        else:
            self.__file = open(self.__file_path, mode, encoding=encoding, *args, **kwargs)
        if is_refresh:
            self.refresh()
        return self.__file
    def close(self):
        if self.__file:
            self.__file.close()
        if self.__datas_lit_key in self.datas:
            self.datas[self.__datas_lit_key] = None
        return self.__file
    def is_open(self)->bool:
        return self.__file is not None

    def load(self):
        if self.__file_path is None:
            if os.path.exists(temp_tool_file_path_name):
                self.data = pickle.load(open(temp_tool_file_path_name, 'rb'))
                return self.data
            else:
                raise FileNotFoundError(f"{self.__file_path} not found, but this ToolFile's target is None")
        elif self.is_dir():
            self.__file = open(os.path.join(self.get_path(), temp_tool_file_path_name), 'rb')
            self.data = pickle.load(self.__file)
            return self.data
        suffix = self.get_extension()
        if suffix == 'json':
            self.load_as_json()
        elif suffix == 'csv':
            self.load_as_csv()
        elif suffix == 'xml':
            self.load_as_xml()
        elif suffix == 'xlsx' or suffix == 'xls':
            self.load_as_excel()
        elif suffix in text_readable_file_type:
            self.load_as_text()
        elif suffix == 'docx' or suffix == 'doc':
            self.load_as_docx()
        elif suffix in audio_file_type:
            self.load_as_audio()
        elif is_image_file(self.__file_path):
            self.load_as_image()
        elif is_binary_file(self.__file_path):
            self.load_as_binary()
        else:
            self.load_as_unknown(suffix)
        return self.data
    def load_as_json(self) -> pd.DataFrame:
        if self.is_open() is False or 'w' in self.__file.mode:
            self.open('r')
        json_data = json.load(self.__file)
        #try:
        #    from pydantic import BaseModel
        #    if "__type" in json_data and "pydantic.BaseModel" in json_data["__type"]:
        #        del json_data["__type"]
        #        json_data = BaseModel.model_validate(json_data)
        #except:
        #    pass
        self.data = json_data
        return self.data
    def load_as_csv(self) -> pd.DataFrame:
        if self.is_open() is False or 'w' in self.__file.mode:
            self.open('r')
        self.data = pd.read_csv(self.__file)
        return self.data
    def load_as_xml(self) -> pd.DataFrame:
        if self.is_open() is False or 'w' in self.__file.mode:
            self.open('r')
        self.data = pd.read_xml(self.__file)
        return self.data
    def load_as_dataframe(self) -> pd.DataFrame:
        if self.is_open() is False or 'w' in self.__file.mode:
            self.open('r')
        self.data = pd.read_csv(self.__file)
        return self.data
    def load_as_excel(self) -> pd.DataFrame:
        if self.is_open() is False or 'w' in self.__file.mode:
            self.open('r')
        self.data = pd.read_excel(self.__file)
        return self.data
    def load_as_binary(self) -> bytes:
        if self.is_open() is False or 'w' in self.__file.mode:
            self.open('rb')
        self.data = self.__file.read()
        return self.data
    def load_as_text(self) -> str:
        if self.is_open() is False or 'w' in self.__file.mode:
            self.open('r')
        self.data = list_byte_to_string(self.__file.readlines())
        return self.data
    def load_as_wav(self):
        self.data = AudioSegment.from_wav(self.__file_path)
        return self.data
    def load_as_audio(self):
        self.data = AudioSegment.from_file(self.__file_path)
        return self.data
    def load_as_image(self) -> ImageFile.ImageFile:
        self.data = Image.open(self.__file_path)
        return self.data
    def load_as_docx(self) -> DocumentObject:
        self.data = Document(self.__file_path)
        return self.data
    def load_as_unknown(self, suffix:str) -> Any:
        return self.load_as_text()
    def load_as_model(self, model:type[BaseModel]) -> BaseModel:
        return model.model_validate(self.load_as_json())

    def save(self, path:Optional[str]=None):
        if path is None and self.__file_path is None:
            raise Exception('No file path specified')
        elif path is None and self.is_dir():
            with open(os.path.join(self.__file_path, temp_tool_file_path_name),'wb') as temp_file:
                pickle.dump(self.data, temp_file)
            return self
        suffix = self.get_extension(path)
        if suffix == 'json':
            self.save_as_json(path)
        elif suffix == 'csv':
            self.save_as_csv(path)
        elif suffix == 'xml':
            self.save_as_xml(path)
        elif suffix == 'xlsx' or suffix == 'xls':
            self.save_as_excel(path)
        elif suffix in text_readable_file_type:
            self.save_as_text(path)
        elif suffix == 'docx':
            self.save_as_docx(path)
        elif suffix in audio_file_type:
            self.save_as_audio(path, suffix)
        elif is_binary_file(self.__file_path):
            self.save_as_binary(path)
        elif is_image_file(self.__file_path):
            self.save_as_image(path)
        else:
            self.save_as_unknown(path)
        return self
    def save_as_json(self, path:Optional[str]=None):
        json_data = self.data
        try:
            from pydantic import BaseModel
            if isinstance(json_data, BaseModel):
                json_data = json_data.model_dump()
                json_data["__type"] = f"{self.data.__class__.__name__}, pydantic.BaseModel"
        except:
            pass
        path = path if path is not None else self.__file_path
        self.close()
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=4)
        return self
    def save_as_csv(self, path:Optional[str]=None):
        path = path if path is not None else self.__file_path
        self.data.to_csv(path)
        return self
    def save_as_xml(self, path:Optional[str]=None):
        path = path if path is not None else self.__file_path
        self.data.to_xml(path)
        return self
    def save_as_dataframe(self, path:Optional[str]=None):
        path = path if path is not None else self.__file_path
        self.data.to_csv(path)
        return self
    def save_as_excel(self, path:Optional[str]=None):
        path = path if path is not None else self.__file_path
        self.data.to_excel(path, index=False)
        return self
    def save_as_binary(self, path:Optional[str]=None):
        if path is not None:
            with open(path, 'wb') as f:
                f.write(self.data)
                f.flush()
        else:
            if self.is_open() is False or 'r' in self.__file.mode:
                self.open('wb')
            self.__file.write(self.data)
            self.__file.flush()
        return self
    def save_as_text(self, path:Optional[str]=None):
        if path is not None:
            with open(path, 'w') as f:
                f.writelines(self.data)
                f.flush()
        else:
            if self.is_open() is False or 'r' in self.__file.mode:
                self.open('w')
            self.__file.writelines(self.data)
            self.__file.flush()
        return self
    def save_as_audio(self, path:Optional[str]=None):
        path = path if path is not None else self.__file_path
        self.data.export(path, format=self.get_extension(path))
        return self
    def save_as_image(self, path:Optional[str]=None):
        path = path if path is not None else self.__file_path
        self.data.save(path)
        return self
    def save_as_docx(self, path:Optional[str]=None):
        if self.data is str:
            self.data = Document()
            table = self.data.add_table(rows=1, cols=1)
            table.cell(0, 0).text = self.data
        path = path if path is not None else self.__file_path
        self.data.save(path)
        return self
    def save_as_unknown(self, path:Optional[str]=None):
        self.save_as_text(path)
    def save_as_model(self, model:type[BaseModel], path:Optional[str]=None):
        self.save_as_json(path)

    def get_size(self) -> int:
        '''
        return:
            return size of directory
        '''
        return os.path.getsize(self.__file_path)
    def get_data_type(self) -> type:
        return type(self.data)
    def has_data_type_is(self, types:Union[type, Sequence[type]]) -> bool:
        if isinstance(types, Sequence) is False:
            return self.get_data_type() == types
        return self.get_data_type() in types
    def get_extension(self, path:str=None):
        if self.is_dir() and path is None:
            raise Exception("Cannot get extension of a directory")
        path = path if path is not None else self.__file_path
        if path is None:
            raise Exception("Cannot get extension without target path")
        return get_extension_name(path)
    def get_path(self):
        return self.get_full_path()
    def get_full_path(self):
        return self.__file_path
    def get_filename(self, is_without_extension = False):
        '''
        if target path is a file, it return filename
        if target path is a directory, it return top directory name
        '''
        if is_without_extension and '.' in self.__file_path:
            return get_base_filename(self.__file_path)[:-(len(self.get_extension())+1)]
        elif self.__file_path[-1] == '/' or self.__file_path[-1] == '\\':
            return get_base_filename(self.__file_path[:-1])
        else:
            return get_base_filename(self.__file_path)
    def get_dir(self):
        return os.path.dirname(self.__file_path)
    def get_dir_tool_file(self):
        return tool_file(self.get_dir())
    def get_current_dir_name(self):
        return os.path.dirname(self.__file_path)

    def is_dir(self):
        if self.__file_path[-1] == '\\' or self.get_path()[-1] == '/':
            return True
        else:
            return os.path.isdir(self.__file_path)
    def is_file(self):
        return os.path.isfile(self.__file_path)
    def is_binary_file(self):
        return is_binary_file(self.__file)
    def is_image(self):
         return is_image_file(self.__file_path)

    def try_create_parent_path(self):
        dir_path = os.path.dirname(self.__file_path)
        if dir_path == '':
            return self
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
        return self
    def dir_iter(self):
        return os.listdir(self.__file_path)
    def dir_tool_file_iter(self):
        result = [self]
        result.clear()
        for file in os.listdir(self.__file_path):
            result.append(self|file)
        return result
    def back_to_parent_dir(self):
        self.close()
        self.__file_path = self.get_dir()
        return self
    def dir_count(self, ignore_folder:bool = True):
        iter    = self.dir_iter()
        result  = 0
        for content in iter:
            if ignore_folder and os.path.isdir(os.path.join(self.__file_path, content)):
                continue
            result += 1
        return result
    def dir_clear(self):
        for file in self.dir_tool_file_iter():
            file.remove()
        return self
    def first_file_with_extension(self, extension:str):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        for file in target_dir.dir_tool_file_iter():
            if file.is_dir() is False and file.get_extension() == extension:
                return file
        return None
    def first_file(self, pr:Callable[[str], bool]):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        for file in target_dir.dir_tool_file_iter():
            if pr(file.get_filename()):
                return file
        return None
    def find_file_with_extension(self, extension:str):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        result:List[tool_file] = []
        for file in target_dir.dir_tool_file_iter():
            if file.is_dir() is False and file.get_extension() == extension:
                result.append(file)
        return result
    def find_file(self, pr:Callable[[str], bool]):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        result:List[tool_file] = []
        for file in target_dir.dir_tool_file_iter():
            if pr(file.get_filename()):
                result.append(file)
        return result
    def dir_walk(
        self,
        top,
        topdown:        bool               = True,
        onerror:        Optional[Callable] = None,
        followlinks:    bool               = False
        ) -> Iterator[tuple[dir_name_type, list[dir_name_type], list[file_name_type]]]:
        return os.walk(self.__file_path, top=top, topdown=topdown, onerror=onerror, followlinks=followlinks)

    def append_text(self, line:str):
        if self.has_data_type_is(type(str)):
            self.data += line
        elif self.has_data_type_is(type(DocumentObject)):
            self.data.add_paragraph(line)
        else:
            raise TypeError(f"Unsupported data type for {sys._getframe().f_code.co_name}")
        return self

    def bool(self):
        return self.exists()
    def __bool__(self):
        return self.exists()

    def must_exists_path(self):
        self.close()
        self.try_create_parent_path()
        self.create()
        return self

    def make_file_inside(self, data:Self, is_delete_source = False):
        if self.is_dir() is False:
            raise Exception("Cannot make file inside a file, because this object target is not a directory")
        result = self|data.get_filename()
        if is_delete_source:
            data.move(result)
        else:
            data.copy(result)
        return self

    @property
    def extension(self):
        if self.is_dir():
            return None
        return self.get_extension()
    @property
    def filename(self):
        if self.is_dir():
            return None
        return self.get_filename(True)
    @property
    def dirname(self):
        if self.is_dir():
            return self.get_current_dir_name()
        return None
    @property
    def dirpath(self):
        if self.is_dir():
            return self.get_current_dir_name()
        return None
    @property
    def shortname(self):
        return self.get_filename(False)
    @property
    def fullpath(self):
        return self.get_path()

    def make_lib_path(self) -> Path:
        return Path(self.__file_path)

    def in_extensions(self, *args:str) -> bool:
        return self.get_extension() in args

    @override
    def SymbolName(self):
        return f"ToolFile<{self.get_path()}>"
    @override
    def ToString(self):
        return self.get_path()

    def compress(self, output_path: Optional[str] = None, format: str = 'zip') -> Self:
        """
        压缩文件或目录
        Args:
            output_path: 输出路径,如果为None则使用原文件名
            format: 压缩格式,支持'zip'和'tar'
        Returns:
            压缩后的文件对象
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        if output_path is None:
            output_path = self.get_path() + ('.zip' if format == 'zip' else '.tar')

        try:
            if format == 'zip':
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    if self.is_dir():
                        for root, _, files in os.walk(self.get_path()):
                            for file in files:
                                file_path = os.path.join(root, file)
                                arcname = os.path.relpath(file_path, self.get_path())
                                zipf.write(file_path, arcname)
                    else:
                        zipf.write(self.get_path(), self.get_filename())
            elif format == 'tar':
                with tarfile.open(output_path, 'w') as tarf:
                    if self.is_dir():
                        tarf.add(self.get_path(), arcname=self.get_filename())
                    else:
                        tarf.add(self.get_path(), arcname=self.get_filename())
            else:
                raise CompressionError(f"Unsupported compression format: {format}")

            return tool_file(output_path)
        except Exception as e:
            raise CompressionError(f"Compression failed: {str(e)}")

    def decompress(self, output_path: Optional[str] = None) -> Self:
        """
        解压文件
        Args:
            output_path: 输出目录,如果为None则使用原文件名
        Returns:
            解压后的目录对象
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        if output_path is None:
            output_path = self.get_path() + '_extracted'

        try:
            if self.get_extension() == 'zip':
                with zipfile.ZipFile(self.get_path(), 'r') as zipf:
                    zipf.extractall(output_path)
            elif self.get_extension() == 'tar':
                with tarfile.open(self.get_path(), 'r') as tarf:
                    tarf.extractall(output_path)
            else:
                raise CompressionError(f"Unsupported archive format: {self.get_extension()}")

            return tool_file(output_path)
        except Exception as e:
            raise CompressionError(f"Decompression failed: {str(e)}")

    def encrypt(self, key: str, algorithm: str = 'AES') -> Self:
        """
        加密文件
        Args:
            key: 加密密钥
            algorithm: 加密算法,目前支持'AES'
        Returns:
            加密后的文件对象
        """
        from cryptography.fernet import Fernet
        from cryptography.hazmat.primitives import hashes
        from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            # 生成加密密钥
            salt = os.urandom(16)
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=salt,
                iterations=100000,
            )
            key = base64.urlsafe_b64encode(kdf.derive(key.encode()))

            # 创建加密器
            f = Fernet(key)

            # 读取文件内容
            with open(self.get_path(), 'rb') as file:
                file_data = file.read()

            # 加密数据
            encrypted_data = f.encrypt(file_data)

            # 保存加密后的文件
            encrypted_path = self.get_path() + '.encrypted'
            with open(encrypted_path, 'wb') as file:
                file.write(salt + encrypted_data)

            return tool_file(encrypted_path)
        except Exception as e:
            raise EncryptionError(f"Encryption failed: {str(e)}")

    def decrypt(self, key: str, algorithm: str = 'AES') -> Self:
        """
        解密文件
        Args:
            key: 解密密钥
            algorithm: 解密算法,目前支持'AES'
        Returns:
            解密后的文件对象
        """
        from cryptography.fernet import Fernet
        from cryptography.hazmat.primitives import hashes
        from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            # 读取加密文件
            with open(self.get_path(), 'rb') as file:
                file_data = file.read()

            # 提取salt和加密数据
            salt = file_data[:16]
            encrypted_data = file_data[16:]

            # 生成解密密钥
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=salt,
                iterations=100000,
            )
            key = base64.urlsafe_b64encode(kdf.derive(key.encode()))

            # 创建解密器
            f = Fernet(key)

            # 解密数据
            decrypted_data = f.decrypt(encrypted_data)

            # 保存解密后的文件
            decrypted_path = self.get_path() + '.decrypted'
            with open(decrypted_path, 'wb') as file:
                file.write(decrypted_data)

            return tool_file(decrypted_path)
        except Exception as e:
            raise EncryptionError(f"Decryption failed: {str(e)}")

    def calculate_hash(self, algorithm: str = 'md5', chunk_size: int = 8192) -> str:
        """
        计算文件的哈希值
        Args:
            algorithm: 哈希算法,支持'md5', 'sha1', 'sha256', 'sha512'等
            chunk_size: 每次读取的字节数
        Returns:
            文件的哈希值(十六进制字符串)
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            # 获取哈希算法
            hash_algo = getattr(hashlib, algorithm.lower())
            if not hash_algo:
                raise HashError(f"Unsupported hash algorithm: {algorithm}")

            # 创建哈希对象
            hasher = hash_algo()

            # 分块读取文件并更新哈希值
            with open(self.get_path(), 'rb') as f:
                while chunk := f.read(chunk_size):
                    hasher.update(chunk)

            return hasher.hexdigest()
        except Exception as e:
            raise HashError(f"Hash calculation failed: {str(e)}")

    def verify_hash(self, expected_hash: str, algorithm: str = 'md5') -> bool:
        """
        验证文件哈希值
        Args:
            expected_hash: 期望的哈希值
            algorithm: 哈希算法,支持'md5', 'sha1', 'sha256', 'sha512'等
        Returns:
            是否匹配
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            actual_hash = self.calculate_hash(algorithm)
            return actual_hash.lower() == expected_hash.lower()
        except Exception as e:
            raise HashError(f"Hash verification failed: {str(e)}")

    def save_hash(self, algorithm: str = 'md5', output_path: Optional[str] = None) -> Self:
        """
        保存文件的哈希值到文件
        Args:
            algorithm: 哈希算法
            output_path: 输出文件路径,如果为None则使用原文件名
        Returns:
            哈希值文件对象
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            # 计算哈希值
            hash_value = self.calculate_hash(algorithm)

            # 生成输出路径
            if output_path is None:
                output_path = self.get_path() + f'.{algorithm}'

            # 保存哈希值
            with open(output_path, 'w') as f:
                f.write(f"{hash_value} *{self.get_filename()}")

            return tool_file(output_path)
        except Exception as e:
            raise HashError(f"Hash saving failed: {str(e)}")

    def start_monitoring(
        self,
        callback:           Callable[[str, str], None],
        recursive:          bool                        = False,
        ignore_patterns:    Optional[List[str]]         = None,
        ignore_directories: bool                        = False,
        case_sensitive:     bool                        = True,
        is_log:             bool                        = True
    ) -> None:
        """
        开始监控文件或目录的变化
        Args:
            callback: 回调函数,接收事件类型和路径两个参数
            recursive: 是否递归监控子目录
            ignore_patterns: 忽略的文件模式列表
            ignore_directories: 是否忽略目录事件
            case_sensitive: 是否区分大小写
        """
        from watchdog.observers import Observer
        from watchdog.events   import FileSystemEventHandler
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            class EventHandler(FileSystemEventHandler):
                def __init__(self, callback, ignore_patterns, ignore_directories, case_sensitive):
                    self.callback = callback
                    self.ignore_patterns = ignore_patterns or []
                    self.ignore_directories = ignore_directories
                    self.case_sensitive = case_sensitive

                def should_ignore(self, path: str) -> bool:
                    if self.ignore_directories and os.path.isdir(path):
                        return True
                    if not self.case_sensitive:
                        path = path.lower()
                    return any(pattern in path for pattern in self.ignore_patterns)

                def on_created(self, event):
                    if not self.should_ignore(event.src_path):
                        self.callback('created', event.src_path)

                def on_modified(self, event):
                    if not self.should_ignore(event.src_path):
                        self.callback('modified', event.src_path)

                def on_deleted(self, event):
                    if not self.should_ignore(event.src_path):
                        self.callback('deleted', event.src_path)

                def on_moved(self, event):
                    if not self.should_ignore(event.src_path):
                        self.callback('moved', f"{event.src_path} -> {event.dest_path}")

            # 创建事件处理器
            event_handler = EventHandler(
                callback=callback,
                ignore_patterns=ignore_patterns,
                ignore_directories=ignore_directories,
                case_sensitive=case_sensitive
            )

            # 创建观察者
            observer = Observer()
            observer.schedule(event_handler, self.get_path(), recursive=recursive)

            # 启动监控
            observer.start()
            if is_log:
                print(f"Started monitoring {self.get_path()}")

            try:
                while True:
                    time.sleep(1)
            except KeyboardInterrupt:
                observer.stop()
                if is_log:
                    print("Stopped monitoring")

            observer.join()

        except Exception as e:
            raise FileMonitorError(f"Failed to start monitoring: {str(e)}")

    def create_backup(
        self,
        backup_dir: Optional[str] = None,
        max_backups: int = 5,
        backup_format: str = 'zip',
        include_metadata: bool = True
    ) -> Self:
        """
        创建文件或目录的备份
        Args:
            backup_dir: 备份目录,如果为None则使用原目录下的.backup目录
            max_backups: 最大保留备份数量
            backup_format: 备份格式,支持'zip'和'tar'
            include_metadata: 是否包含元数据
        Returns:
            备份文件对象
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            # 生成备份目录
            if backup_dir is None:
                backup_dir = os.path.join(self.get_dir(), '.backup')
            backup_dir:Self = tool_file(backup_dir)
            backup_dir.must_exists_path()

            # 生成备份文件名
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_name = f"{self.get_filename()}_{timestamp}"

            # 创建备份
            if backup_format == 'zip':
                backup_path = backup_dir | f"{backup_name}.zip"
                with zipfile.ZipFile(backup_path.get_path(), 'w', zipfile.ZIP_DEFLATED) as zipf:
                    if self.is_dir():
                        for root, _, files in os.walk(self.get_path()):
                            for file in files:
                                file_path = os.path.join(root, file)
                                arcname = os.path.relpath(file_path, self.get_path())
                                zipf.write(file_path, arcname)
                    else:
                        zipf.write(self.get_path(), self.get_filename())
            elif backup_format == 'tar':
                backup_path = backup_dir | f"{backup_name}.tar"
                with tarfile.open(backup_path.get_path(), 'w') as tarf:
                    if self.is_dir():
                        tarf.add(self.get_path(), arcname=self.get_filename())
                    else:
                        tarf.add(self.get_path(), arcname=self.get_filename())
            else:
                raise BackupError(f"Unsupported backup format: {backup_format}")

            # 添加元数据
            if include_metadata:
                metadata = {
                    'original_path': self.get_path(),
                    'backup_time': timestamp,
                    'file_size': self.get_size(),
                    'is_directory': self.is_dir(),
                    'hash': self.calculate_hash()
                }
                metadata_path = backup_dir | f"{backup_name}.meta.json"
                with open(metadata_path.get_path(), 'w') as f:
                    json.dump(metadata, f, indent=4)

            # 清理旧备份
            if max_backups > 0:
                backups = backup_dir.find_file(lambda f: tool_file(f).get_filename().startswith(self.get_filename() + '_'))
                backups.sort(key=lambda f: f.get_filename(), reverse=True)
                for old_backup in backups[max_backups:]:
                    old_backup.remove()

            return backup_path

        except Exception as e:
            raise BackupError(f"Backup failed: {str(e)}")

    def restore_backup(
        self,
        backup_file: Union[str, Self],
        restore_path: Optional[str] = None,
        verify_hash: bool = True
    ) -> Self:
        """
        从备份恢复文件或目录
        Args:
            backup_file: 备份文件路径
            restore_path: 恢复路径,如果为None则恢复到原位置
            verify_hash: 是否验证哈希值
        Returns:
            恢复后的文件对象
        """
        if not isinstance(backup_file, tool_file):
            backup_file:Self = tool_file(backup_file)

        if not backup_file.exists():
            raise FileNotFoundError(f"Backup file not found: {backup_file.get_path()}")

        try:
            # 确定恢复路径
            if restore_path is None:
                restore_path = self.get_path()
            restore_path:Self = tool_file(restore_path)

            # 解压备份
            if backup_file.get_extension() == 'zip':
                with zipfile.ZipFile(backup_file.get_path(), 'r') as zipf:
                    zipf.extractall(restore_path.get_path())
            elif backup_file.get_extension() == 'tar':
                with tarfile.open(backup_file.get_path(), 'r') as tarf:
                    tarf.extractall(restore_path.get_path())
            else:
                raise BackupError(f"Unsupported backup format: {backup_file.get_extension()}")

            # 验证哈希值
            if verify_hash:
                metadata_path = backup_file.get_path()[:-len(backup_file.get_extension())-1] + '.meta.json'
                if os.path.exists(metadata_path):
                    with open(metadata_path, 'r') as f:
                        metadata = json.load(f)
                    restored_file = tool_file(restore_path.get_path())
                    if restored_file.calculate_hash() != metadata['hash']:
                        raise BackupError("Hash verification failed")

            return restore_path

        except Exception as e:
            raise BackupError(f"Restore failed: {str(e)}")

    def list_backups(self) -> List[Self]:
        """
        列出所有备份
        Returns:
            备份文件列表
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            backup_dir:Self = tool_file(os.path.join(self.get_dir(), '.backup'))
            if not backup_dir.exists():
                return []

            backups = backup_dir.find_file(lambda f: tool_file(f).get_filename().startswith(self.get_filename() + '_'))
            backups.sort(key=lambda f: tool_file(f).get_filename(), reverse=True)
            return backups

        except Exception as e:
            raise BackupError(f"Failed to list backups: {str(e)}")

    def get_permissions(self) -> Dict[str, bool]:
        """
        获取文件或目录的权限
        Returns:
            权限字典,包含以下键:
            - read: 是否可读
            - write: 是否可写
            - execute: 是否可执行
            - hidden: 是否隐藏
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            mode = os.stat(self.get_path()).st_mode
            return {
                'read': bool(mode & stat.S_IRUSR),
                'write': bool(mode & stat.S_IWUSR),
                'execute': bool(mode & stat.S_IXUSR),
                'hidden': bool(os.path.isfile(self.get_path()) and self.get_filename().startswith('.'))
            }
        except Exception as e:
            raise PermissionError(f"Failed to get permissions: {str(e)}")

    def set_permissions(
        self,
        read: Optional[bool] = None,
        write: Optional[bool] = None,
        execute: Optional[bool] = None,
        hidden: Optional[bool] = None,
        recursive: bool = False
    ) -> Self:
        """
        设置文件或目录的权限
        Args:
            read: 是否可读
            write: 是否可写
            execute: 是否可执行
            hidden: 是否隐藏
            recursive: 是否递归设置目录权限
        Returns:
            文件对象本身
        """
        if not self.exists():
            raise FileNotFoundError(f"File not found: {self.get_path()}")

        try:
            # 获取当前权限
            current_perms = os.stat(self.get_path()).st_mode

            # 设置新权限
            if read is not None:
                if read:
                    current_perms |= stat.S_IRUSR
                else:
                    current_perms &= ~stat.S_IRUSR

            if write is not None:
                if write:
                    current_perms |= stat.S_IWUSR
                else:
                    current_perms &= ~stat.S_IWUSR

            if execute is not None:
                if execute:
                    current_perms |= stat.S_IXUSR
                else:
                    current_perms &= ~stat.S_IXUSR

            # 应用权限
            os.chmod(self.get_path(), current_perms)

            # 设置隐藏属性
            if hidden is not None:
                if os.name == 'nt':  # Windows
                    import ctypes
                    if hidden:
                        ctypes.windll.kernel32.SetFileAttributesW(self.get_path(), 2)
                    else:
                        ctypes.windll.kernel32.SetFileAttributesW(self.get_path(), 0)
                else:  # Unix/Linux/Mac
                    if hidden:
                        if not self.get_filename().startswith('.'):
                            self.rename('.' + self.get_filename())
                    else:
                        if self.get_filename().startswith('.'):
                            self.rename(self.get_filename()[1:])

            # 递归设置目录权限
            if recursive and self.is_dir():
                for root, _, files in os.walk(self.get_path()):
                    for file in files:
                        file_path = os.path.join(root, file)
                        if read is not None:
                            if read:
                                os.chmod(file_path, os.stat(file_path).st_mode | stat.S_IRUSR)
                            else:
                                os.chmod(file_path, os.stat(file_path).st_mode & ~stat.S_IRUSR)
                        if write is not None:
                            if write:
                                os.chmod(file_path, os.stat(file_path).st_mode | stat.S_IWUSR)
                            else:
                                os.chmod(file_path, os.stat(file_path).st_mode & ~stat.S_IWUSR)
                        if execute is not None:
                            if execute:
                                os.chmod(file_path, os.stat(file_path).st_mode | stat.S_IXUSR)
                            else:
                                os.chmod(file_path, os.stat(file_path).st_mode & ~stat.S_IXUSR)

            return self

        except Exception as e:
            raise PermissionError(f"Failed to set permissions: {str(e)}")

    def is_readable(self) -> bool:
        """
        检查文件是否可读
        Returns:
            是否可读
        """
        return self.get_permissions()['read']

    def is_writable(self) -> bool:
        """
        检查文件是否可写
        Returns:
            是否可写
        """
        return self.get_permissions()['write']

    def is_executable(self) -> bool:
        """
        检查文件是否可执行
        Returns:
            是否可执行
        """
        return self.get_permissions()['execute']

    def is_hidden(self) -> bool:
        """
        检查文件是否隐藏
        Returns:
            是否隐藏
        """
        return self.get_permissions()['hidden']

class loss_file(tool_file):
    def __init__(self, file_path:str, *args, **kwargs):
        super().__init__(file_path, *args, **kwargs)

    @override
    def __setitem__(self, key:str, value):
        raise ValueError("loss file cannt set item in its data-items")
    @override
    def __getitem__(self, key:str):
        return None
    def __contains__(self, key:str):
        return False
    def __delitem__(self, key:str):
        pass
    def __iter__(self):
        return iter({})
    def __len__(self):
        return -1
    @override
    def __enter__(self):
        return self
    @override
    def __exit__(self, exc_type, exc_val, exc_tb):
        return True

    @override
    def __or__(self, other):
        return loss_file(super().__or__(other).get_path())

    @override
    def create(self) -> tool_file:
        return self
    @override
    def exists(self) -> bool:
        return False
    @override
    def remove(self) -> tool_file:
        return self
    @override
    def copy(self, to_path:Union[Self, str]) -> tool_file:
        return None
    @override
    def move(self, to_path:Union[Self, str]) -> tool_file:
        return self
    @override
    def rename(self, newpath:Union[Self, str]) -> tool_file:
        return self

    override
    def refresh(self) -> tool_file:
        return self
    @override
    def open(self, *args, **kwargs) -> IO[Any]:
        return None
    @override
    def close(self) -> IO[Any]:
        return None
    @override
    def is_open(self)->bool:
        return False

    @override
    def load(self):
        return None
    @override
    def load_as_json(self) -> pd.DataFrame:
        return None
    @override
    def load_as_csv(self) -> pd.DataFrame:
        return None
    @override
    def load_as_xml(self) -> pd.DataFrame:
        return None
    @override
    def load_as_dataframe(self) -> pd.DataFrame:
        return None
    @override
    def load_as_excel(self) -> pd.DataFrame:
        return None
    @override
    def load_as_binary(self) -> bytes:
        return None
    @override
    def load_as_text(self) -> str:
        return None
    @override
    def load_as_wav(self):
        return None
    @override
    def load_as_audio(self):
        return None
    @override
    def load_as_image(self) -> ImageFile.ImageFile:
        return None
    @override
    def load_as_docx(self) -> DocumentObject:
        return None
    @override
    def load_as_unknown(self, suffix:str) -> Any:
        return None

    @override
    def save(self, path:str=None):
        return self
    @override
    def save_as_json(self, path:str):
        return self
    @override
    def save_as_csv(self, path:str):
        return self
    @override
    def save_as_xml(self, path:str):
        return self
    @override
    def save_as_dataframe(self, path:str):
        return self
    @override
    def save_as_excel(self, path:str):
        return self
    @override
    def save_as_binary(self, path:str):
        return self
    @override
    def save_as_text(self, path:str):
        return self
    @override
    def save_as_audio(self, path:str):
        return Self
    @override
    def save_as_image(self, path:str):
        return Self
    @override
    def save_as_docx(self, path:str):
        return Self
    @override
    def save_as_unknown(self, path:str):
        return Self

    @override
    def compress(self, output_path: Optional[str] = None, format: str = 'zip') -> Self:
        return self

    @override
    def decompress(self, output_path: Optional[str] = None) -> Self:
        return self

    @override
    def encrypt(self, key: str, algorithm: str = 'AES') -> Self:
        return self

    @override
    def decrypt(self, key: str, algorithm: str = 'AES') -> Self:
        return self

    @override
    def calculate_hash(self, algorithm: str = 'md5', chunk_size: int = 8192) -> str:
        return ""

    @override
    def verify_hash(self, expected_hash: str, algorithm: str = 'md5') -> bool:
        return False

    @override
    def save_hash(self, algorithm: str = 'md5', output_path: Optional[str] = None) -> Self:
        return self

    @override
    def SymbolName(self):
        return f"LossFile<{self.get_path()}>"

    @override
    def start_monitoring(
        self,
        callback: Callable[[str, str], None],
        recursive: bool = False,
        ignore_patterns: Optional[List[str]] = None,
        ignore_directories: bool = False,
        case_sensitive: bool = True
    ) -> None:
        pass

    @override
    def stop_monitoring(self) -> None:
        pass

    @override
    def create_backup(
        self,
        backup_dir: Optional[str] = None,
        max_backups: int = 5,
        backup_format: str = 'zip',
        include_metadata: bool = True
    ) -> Self:
        return self

    @override
    def restore_backup(
        self,
        backup_file: Union[str, Self],
        restore_path: Optional[str] = None,
        verify_hash: bool = True
    ) -> Self:
        return self

    @override
    def list_backups(self) -> List[Self]:
        return []

    @override
    def get_permissions(self) -> Dict[str, bool]:
        return {'read': False, 'write': False, 'execute': False, 'hidden': False}

    @override
    def set_permissions(
        self,
        read: Optional[bool] = None,
        write: Optional[bool] = None,
        execute: Optional[bool] = None,
        hidden: Optional[bool] = None,
        recursive: bool = False
    ) -> Self:
        return self

    @override
    def is_readable(self) -> bool:
        return False

    @override
    def is_writable(self) -> bool:
        return False

    @override
    def is_executable(self) -> bool:
        return False

    @override
    def is_hidden(self) -> bool:
        return False

static_loss_file_dir = loss_file(".temp/")
static_loss_file = loss_file(".temp.bad")

def is_loss_tool_file(from_) -> bool:
    return isinstance(from_, loss_file)

def Wrapper(file) -> tool_file:
    if isinstance(file, tool_file):
        return file
    else:
        return tool_file(UnWrapper(file))

def split_elements(
    file:               Union[tool_file, str],
    *,
    ratios:             List[float]                                 = [1,1],
    pr:                 Optional[Callable[[tool_file], bool]]       = None,
    shuffler:           Optional[Callable[[List[tool_file]], None]] = None,
    output_dirs:        Optional[List[tool_file]]                   = None,
    output_must_exist:  bool                                        = True,
    output_callback:    Optional[Callable[[tool_file], None]]       = None
    ) -> List[List[tool_file]]:
    if is_loss_tool_file(file):
        return []
    result:                 List[List[tool_file]]   = tool_split_elements(Wrapper(file).dir_tool_file_iter(),
                                      ratios=ratios,
                                      pr=pr,
                                      shuffler=shuffler)
    if output_dirs is None:
        return result
    for i in range(min(len(output_dirs), len(result))):
        output_dir:         tool_file               = output_dirs[i]
        if output_dir.is_dir() is False:
            raise Exception("Outputs must be directory")
        if output_must_exist:
            output_dir.must_exists_as_new()
        for file in result[i]:
            current = output_dirs[i].make_file_inside(file)
            if output_callback:
                output_callback(current)

    return result

tool_file_or_str = Union[tool_file, str]

if __name__ == "__main__":
    a = tool_file("abc/")
    b = tool_file("a.x")
    c = a|b
    print(c.get_path())


